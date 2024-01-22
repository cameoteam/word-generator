from docx import Document


def is_similar_word_files(
    first_path,
    second_path,
):
    test_file = Document(first_path)
    result_file = Document(second_path)
    for paragraph in range(len(test_file.paragraphs)):
        if test_file.paragraphs[paragraph].text != result_file.paragraphs[paragraph].text:
            return False
    return True
