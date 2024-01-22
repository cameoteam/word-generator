from docx import Document
from .types import InsertedData
from typing import List


class WordManager:
    """Manage word files.

    Parameters
    ----------
    template_path : str
        Path to word file.

    """

    def __init__(self, template_path: str):
        self.document = Document(template_path)

    def save(self, path: str):
        """Save document.

        Parameters
        ----------
        path : str
            The path where the file will be saved.

        """

        self.document.save(path)

    def insert_values(
        self,
        fields: List[InsertedData],
    ):
        """Insert values inside document.

        Parameters
        ----------
        fields : List[InsertedData]
            List of data.

        """

        map(self.insert_value, fields)

    def insert_value(
        self,
        inserted_data: InsertedData,
    ):
        """Insert single value inside document.

        Parameters
        ----------
        inserted_data: InsertedData
            Data to inserted in file.

        """

        key = "{}{}{}".format("{", inserted_data.key, "}")
        for paragraph in self.get_all_paragraphs():
            runs = paragraph.runs
            for run in runs:
                if key not in run.text:
                    continue
                len_text = len(run.text)
                if inserted_data.category == "picture":
                    run.text = ""
                    run.add_picture(inserted_data.value)
                else:
                    run_text = run.text.replace(str(key), inserted_data.value)
                    run.text = self.reduce_line(run_text, len_text)

    def get_all_paragraphs(self) -> list:
        """Get all paragraphs of document."""

        paragraphs = self.document.paragraphs
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs += cell.paragraphs
        return paragraphs

    def reduce_line(self, text: str, length: int) -> str:
        """Delete all unused spaces of line.

        Parameters
        ----------
        text : str
            Line.
        length : int
            The length to be trimmed to.

        """

        while True:
            try:
                if text[-1] == " " and len(text) == length:
                    text = text[:-1]
                else:
                    break
            except IndexError:
                break
        return text
